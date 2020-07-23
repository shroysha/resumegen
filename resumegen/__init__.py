from docx import Document


class ResumeGenerator:

    def __init__(self, save_location):
        self._save_location = save_location
        self._document = Document()

    def __enter__(self):
        self._document.save(self._save_location)

    def __exit__(self, exc_type, exc_val, exc_tb):
        pass


def add_contact_info_box(document, contact_info):
    contact_info_box = document.add_table(rows=2, cols=2)
    contact_info_box.cell(0, 0).text = contact_info.name
    contact_info_box.cell(0, 1).text = contact_info.phone
    contact_info_box.cell(1, 0).text = contact_info.email
    contact_info_box.cell(1, 1).text = contact_info.website


def add_all_certifications(document, certifications):
    document.add_heading("Certifications")
    cert_table = document.add_table(rows=len(certifications), cols=3)
    for cert, cert_row_ind in zip(certifications, range(len(certifications))):
        cert_row = cert_table.row_cells(cert_row_ind)
        cert_row[0].text = cert.acronym
        cert_row[1].text = cert.name
        cert_row[2].text = cert.start_date


def add_all_educations(document, educations):
    for education in educations:
        pass


def add_all_work_experiences(document, work_experiences):
    for work_experience in work_experiences:
        document.add_table(rows=2, cols=2)
        pass


class OriginalResumeGenerator(ResumeGenerator):
    TITLE_TEXT = "Shawn Shroyer's Resume"
    TITLE_HEADING_LEVEL = 0

    def __init__(self, resume_info, save_location):
        super().__init__(save_location)

        document = self._document

        add_contact_info_box(document, resume_info.contact_info)
        add_all_certifications(document, resume_info.certifications)
        add_all_educations(document, resume_info.educations)
        add_all_work_experiences(document, resume_info.work_experiences)

        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='Intense Quote')

        document.add_paragraph(
            'first item in unordered list', style='List Bullet'
        )
        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )

        # document.add_picture('monty-truth.png', width=Inches(1.25))

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        document.add_page_break()


class IndeedResumeGenerator(ResumeGenerator):

    def __init__(self, resume_info, save_location):
        super().__init__(save_location)


def generate_resume(resume_info, resume_format, save_location):
    with resume_format(resume_info) as doc_data:
        pass
    pass
